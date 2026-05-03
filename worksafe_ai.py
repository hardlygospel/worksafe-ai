#!/usr/bin/env python3
# SPDX-License-Identifier: GPL-3.0-or-later
#
# Worksafe AI — Private Local LLM for home use
# Copyright (C) 2024  hardlygospel
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.
"""
Worksafe AI — Private AI for home use.
Powered by Ollama · 100% Offline · No API Keys Required.
"""
from __future__ import annotations

import argparse
import base64
import json
import platform
import re
import shutil
import subprocess
import sys
import time
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional


def _pip_install(packages: list[str]) -> None:
    """Install packages, trying --user if the environment is externally managed."""
    for extra in ([], ["--user"], ["--user", "--break-system-packages"]):
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", "--quiet"] + extra + packages,
                stderr=subprocess.DEVNULL,
            )
            return
        except subprocess.CalledProcessError:
            continue
    print(f"\n[!] Could not auto-install: {' '.join(packages)}")
    print(f"    Please run:  pip install {' '.join(packages)}\n")
    sys.exit(1)


def _ensure_deps() -> None:
    missing = []
    try:
        import rich  # noqa: F401
    except ImportError:
        missing.append("rich>=13.7.0")
    try:
        import requests  # noqa: F401
    except ImportError:
        missing.append("requests>=2.31.0")
    if missing:
        _pip_install(missing)


_ensure_deps()

import requests  # noqa: E402
from rich import box  # noqa: E402
from rich.console import Console  # noqa: E402
from rich.markdown import Markdown  # noqa: E402
from rich.padding import Padding  # noqa: E402
from rich.panel import Panel  # noqa: E402
from rich.prompt import Confirm, Prompt  # noqa: E402
from rich.rule import Rule  # noqa: E402
from rich.table import Table  # noqa: E402
from rich.text import Text  # noqa: E402

console = Console()

# ── Constants ─────────────────────────────────────────────────────────────────
OLLAMA_API  = "http://localhost:11434"
APP_VERSION   = "3.0.0"
SESSIONS_DIR  = Path.home() / ".worksafe_ai" / "sessions"
VISION_MODELS = frozenset({"llava", "moondream", "bakllava", "llava-llama3", "llava-phi3"})

DEFAULT_SYSTEM_PROMPT = (
    "You are Worksafe AI, a private assistant running entirely on the user's local machine. "
    "No data is ever transmitted to any external server. Be concise, clear, and genuinely "
    "useful. If asked about privacy or data handling, reassure the user that everything stays "
    "on their own hardware — nothing is logged, tracked, or shared."
)

# ── Model catalogue ───────────────────────────────────────────────────────────
MODELS: list[dict] = [
    # ── Fast & Light ──────────────────────────────────────────────────────────
    {
        "id": "llama3.2:1b",    "name": "Llama 3.2 1B",
        "desc": "Ultra-fast, basic tasks & quick questions",
        "size": "~0.7 GB", "min_ram": 2,
        "category": "⚡  Fast & Light", "color": "bright_green",
    },
    {
        "id": "llama3.2:3b",    "name": "Llama 3.2 3B",
        "desc": "Fast everyday assistant — great starting point",
        "size": "~2 GB",   "min_ram": 4,
        "category": "⚡  Fast & Light", "color": "bright_green", "recommended": True,
    },
    {
        "id": "gemma3:1b",      "name": "Gemma 3 1B",
        "desc": "Google's smallest capable model",
        "size": "~0.8 GB", "min_ram": 2,
        "category": "⚡  Fast & Light", "color": "green",
    },
    {
        "id": "phi3.5:mini",    "name": "Phi 3.5 Mini",
        "desc": "Microsoft — efficient & capable for its size",
        "size": "~2.3 GB", "min_ram": 4,
        "category": "⚡  Fast & Light", "color": "green",
    },
    # ── Balanced ──────────────────────────────────────────────────────────────
    {
        "id": "llama3.1:8b",    "name": "Llama 3.1 8B",
        "desc": "Best everyday balance of speed & quality",
        "size": "~5 GB",   "min_ram": 8,
        "category": "⚖️   Balanced",    "color": "cyan", "recommended": True,
    },
    {
        "id": "mistral:7b",     "name": "Mistral 7B",
        "desc": "Excellent reasoning, writing & summarisation",
        "size": "~4 GB",   "min_ram": 8,
        "category": "⚖️   Balanced",    "color": "cyan",
    },
    {
        "id": "gemma3:4b",      "name": "Gemma 3 4B",
        "desc": "Compact & capable — great on Apple Silicon",
        "size": "~3 GB",   "min_ram": 6,
        "category": "⚖️   Balanced",    "color": "blue",
    },
    {
        "id": "gemma3:12b",     "name": "Gemma 3 12B",
        "desc": "Excellent quality general-purpose model",
        "size": "~8 GB",   "min_ram": 12,
        "category": "⚖️   Balanced",    "color": "blue",
    },
    {
        "id": "qwen2.5:7b",     "name": "Qwen 2.5 7B",
        "desc": "Strong multilingual support (29 languages)",
        "size": "~5 GB",   "min_ram": 8,
        "category": "⚖️   Balanced",    "color": "bright_cyan",
    },
    # ── Reasoning ─────────────────────────────────────────────────────────────
    {
        "id": "deepseek-r1:8b",  "name": "DeepSeek R1 8B",
        "desc": "Chain-of-thought reasoning & analysis",
        "size": "~5 GB",   "min_ram": 8,
        "category": "🧠  Reasoning",    "color": "bright_magenta",
    },
    {
        "id": "deepseek-r1:14b", "name": "DeepSeek R1 14B",
        "desc": "Strong multi-step reasoning",
        "size": "~9 GB",   "min_ram": 12,
        "category": "🧠  Reasoning",    "color": "magenta",
    },
    {
        "id": "phi4:14b",        "name": "Phi 4 14B",
        "desc": "Microsoft's best compact reasoning model",
        "size": "~8 GB",   "min_ram": 12,
        "category": "🧠  Reasoning",    "color": "bright_magenta",
    },
    {
        "id": "qwq:32b",         "name": "QwQ 32B",
        "desc": "Frontier-level reasoning (needs 32 GB+ RAM)",
        "size": "~20 GB",  "min_ram": 32,
        "category": "🧠  Reasoning",    "color": "magenta",
    },
    # ── Coding ────────────────────────────────────────────────────────────────
    {
        "id": "codellama:7b",          "name": "CodeLlama 7B",
        "desc": "Meta's code specialist — all languages",
        "size": "~4 GB",   "min_ram": 8,
        "category": "💻  Coding",       "color": "bright_blue",
    },
    {
        "id": "codellama:13b",         "name": "CodeLlama 13B",
        "desc": "Better code generation & explanations",
        "size": "~8 GB",   "min_ram": 12,
        "category": "💻  Coding",       "color": "blue",
    },
    {
        "id": "codegemma:7b",          "name": "CodeGemma 7B",
        "desc": "Google's coding model, strong Python & JS",
        "size": "~5 GB",   "min_ram": 8,
        "category": "💻  Coding",       "color": "bright_blue",
    },
    {
        "id": "deepseek-coder-v2:16b", "name": "DeepSeek Coder V2 16B",
        "desc": "Top-tier code generation & debugging",
        "size": "~10 GB",  "min_ram": 14,
        "category": "💻  Coding",       "color": "blue",
    },
    {
        "id": "starcoder2:15b",        "name": "StarCoder 2 15B",
        "desc": "Trained on 600+ programming languages",
        "size": "~9 GB",   "min_ram": 12,
        "category": "💻  Coding",       "color": "bright_blue",
    },
    # ── Multilingual ──────────────────────────────────────────────────────────
    {
        "id": "qwen2.5:14b",    "name": "Qwen 2.5 14B",
        "desc": "High-quality in 29 languages",
        "size": "~9 GB",   "min_ram": 12,
        "category": "🌍  Multilingual", "color": "bright_cyan",
    },
    {
        "id": "aya:8b",         "name": "Aya 8B",
        "desc": "Cohere's model tuned for 23 languages",
        "size": "~5 GB",   "min_ram": 8,
        "category": "🌍  Multilingual", "color": "cyan",
    },
    {
        "id": "aya-expanse:8b", "name": "Aya Expanse 8B",
        "desc": "Improved multilingual instruction following",
        "size": "~5 GB",   "min_ram": 8,
        "category": "🌍  Multilingual", "color": "cyan",
    },
    # ── High Power ────────────────────────────────────────────────────────────
    {
        "id": "mixtral:8x7b",   "name": "Mixtral 8x7B",
        "desc": "Mixture-of-experts architecture, very capable",
        "size": "~26 GB",  "min_ram": 32,
        "category": "🔋  High Power",   "color": "yellow",
    },
    {
        "id": "qwen2.5:32b",    "name": "Qwen 2.5 32B",
        "desc": "Excellent large multilingual model",
        "size": "~20 GB",  "min_ram": 24,
        "category": "🔋  High Power",   "color": "bright_yellow",
    },
    {
        "id": "llama3.1:70b",   "name": "Llama 3.1 70B",
        "desc": "Meta's largest open model — near GPT-4 quality",
        "size": "~40 GB",  "min_ram": 48,
        "category": "🔋  High Power",   "color": "bright_yellow",
    },
    {
        "id": "deepseek-r1:70b","name": "DeepSeek R1 70B",
        "desc": "Frontier-class chain-of-thought reasoning",
        "size": "~40 GB",  "min_ram": 48,
        "category": "🔋  High Power",   "color": "yellow",
    },
    {
        "id": "gemma3:27b",     "name": "Gemma 3 27B",
        "desc": "Google's largest Gemma — near GPT-4 quality",
        "size": "~17 GB",  "min_ram": 24,
        "category": "🔋  High Power",   "color": "bright_yellow",
    },
    {
        "id": "command-r:35b",  "name": "Command R 35B",
        "desc": "Cohere — optimised for RAG & tool use",
        "size": "~20 GB",  "min_ram": 24,
        "category": "🔋  High Power",   "color": "yellow",
    },
    # ── Balanced additions ────────────────────────────────────────────────────
    {
        "id": "mistral-nemo:12b","name": "Mistral Nemo 12B",
        "desc": "Mistral's newer efficient 12B — very strong",
        "size": "~7 GB",   "min_ram": 10,
        "category": "⚖️   Balanced",    "color": "bright_cyan",
    },
    {
        "id": "llama3.2:11b",   "name": "Llama 3.2 11B",
        "desc": "Meta's vision-capable text + image model",
        "size": "~8 GB",   "min_ram": 12,
        "category": "⚖️   Balanced",    "color": "cyan",
    },
    # ── Coding additions ──────────────────────────────────────────────────────
    {
        "id": "qwen2.5-coder:7b","name": "Qwen 2.5 Coder 7B",
        "desc": "Alibaba's code-specific model — strong completions",
        "size": "~5 GB",   "min_ram": 8,
        "category": "💻  Coding",       "color": "bright_blue",
    },
    {
        "id": "qwen2.5-coder:14b","name": "Qwen 2.5 Coder 14B",
        "desc": "Larger Qwen coder — excellent for complex tasks",
        "size": "~9 GB",   "min_ram": 12,
        "category": "💻  Coding",       "color": "blue",
    },
    # ── Vision & Multimodal ───────────────────────────────────────────────────
    {
        "id": "llava:7b",       "name": "LLaVA 7B",
        "desc": "Vision-language — describe & reason about images",
        "size": "~5 GB",   "min_ram": 8,
        "category": "👁️   Vision",      "color": "bright_yellow",
    },
    {
        "id": "llava:13b",      "name": "LLaVA 13B",
        "desc": "Better image understanding & analysis",
        "size": "~8 GB",   "min_ram": 12,
        "category": "👁️   Vision",      "color": "yellow",
    },
    {
        "id": "moondream:1.8b", "name": "Moondream 1.8B",
        "desc": "Tiny vision model — fast image Q&A on low-end hardware",
        "size": "~1.7 GB", "min_ram": 4,
        "category": "👁️   Vision",      "color": "bright_yellow",
    },
    {
        "id": "llava-llama3:8b","name": "LLaVA Llama 3 8B",
        "desc": "LLaVA on Llama 3 backbone — strong visual reasoning",
        "size": "~5 GB",   "min_ram": 8,
        "category": "👁️   Vision",      "color": "yellow",
    },
]

MODEL_LOOKUP: dict[str, dict] = {m["id"]: m for m in MODELS}


def _is_installed(model: str, local: list[str]) -> bool:
    """Exact match, then untagged base-name match (e.g. 'mistral' → 'mistral:7b')."""
    if model in local:
        return True
    if ":" not in model:
        return any(m.split(":")[0] == model for m in local)
    return False


def _is_vision_model(model: str) -> bool:
    base = model.split(":")[0].lower()
    return any(v in base for v in VISION_MODELS)


# ── Hardware detection ────────────────────────────────────────────────────────

@dataclass
class HardwareInfo:
    gpu_type:  str = "none"
    gpu_name:  str = "No dedicated GPU detected"
    vram_gb:   Optional[int] = None
    ram_gb:    int = 8
    recommended: list[str] = field(default_factory=list)


def _system_ram_gb() -> int:
    try:
        sys_name = platform.system()
        if sys_name == "Darwin":
            out = subprocess.run(
                ["sysctl", "-n", "hw.memsize"], capture_output=True, text=True, timeout=5
            )
            return int(out.stdout.strip()) // (1024 ** 3)
        if sys_name == "Linux":
            with open("/proc/meminfo") as f:
                for line in f:
                    if line.startswith("MemTotal"):
                        return int(line.split()[1]) // (1024 * 1024)
        if sys_name == "Windows":
            out = subprocess.run(
                ["powershell", "-Command",
                 "(Get-CimInstance Win32_ComputerSystem).TotalPhysicalMemory"],
                capture_output=True, text=True, timeout=5,
            )
            return int(out.stdout.strip()) // (1024 ** 3)
    except Exception:
        pass
    return 8


def detect_hardware() -> HardwareInfo:
    hw = HardwareInfo(ram_gb=_system_ram_gb())

    # Apple Silicon — unified memory acts as both RAM and VRAM
    if platform.system() == "Darwin" and platform.machine() == "arm64":
        try:
            out = subprocess.run(
                ["sysctl", "-n", "machdep.cpu.brand_string"],
                capture_output=True, text=True, timeout=5,
            )
            cpu = out.stdout.strip()
            if "Apple" in cpu:
                hw.gpu_type = "apple_silicon"
                hw.gpu_name = cpu
                hw.vram_gb  = hw.ram_gb  # unified memory
        except Exception:
            pass

    # NVIDIA
    if hw.gpu_type == "none":
        try:
            out = subprocess.run(
                ["nvidia-smi",
                 "--query-gpu=name,memory.total",
                 "--format=csv,noheader,nounits"],
                capture_output=True, text=True, timeout=5,
            )
            if out.returncode == 0 and out.stdout.strip():
                parts = out.stdout.strip().split(", ")
                hw.gpu_type = "nvidia"
                hw.gpu_name = parts[0].strip()
                hw.vram_gb  = int(parts[1].strip()) // 1024
        except (FileNotFoundError, subprocess.TimeoutExpired, ValueError):
            pass

    # AMD (ROCm on Linux)
    if hw.gpu_type == "none" and platform.system() == "Linux":
        try:
            out = subprocess.run(
                ["rocm-smi", "--showproductname"],
                capture_output=True, text=True, timeout=5,
            )
            if out.returncode == 0 and out.stdout.strip():
                hw.gpu_type = "amd"
                hw.gpu_name = out.stdout.strip().split("\n")[0].strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

    # Build recommendations
    effective = hw.vram_gb if hw.vram_gb else hw.ram_gb
    # For CPU-only, models need ~1.5× their size in RAM — be conservative
    limit = effective if hw.gpu_type != "none" else int(effective * 0.6)
    candidates = [m["id"] for m in MODELS if m["min_ram"] <= limit]
    # Sort: recommended flag first, then largest that still fits
    candidates.sort(key=lambda x: (
        not MODEL_LOOKUP[x].get("recommended", False),
        MODEL_LOOKUP[x]["min_ram"],
    ))
    hw.recommended = candidates[:5]
    return hw


def show_hardware_panel(hw: HardwareInfo) -> None:
    gpu_icon = {"apple_silicon": "🍎", "nvidia": "🟢", "amd": "🔴"}.get(hw.gpu_type, "🖥️")
    lines: list[str] = []

    if hw.gpu_type == "apple_silicon":
        lines.append(f"{gpu_icon}  [bold bright_cyan]{hw.gpu_name}[/bold bright_cyan]")
        lines.append(f"   Unified Memory: [bold]{hw.ram_gb} GB[/bold]  "
                     f"[dim](shared CPU + GPU)[/dim]")
    elif hw.gpu_type == "nvidia":
        lines.append(f"{gpu_icon}  [bold bright_green]NVIDIA {hw.gpu_name}[/bold bright_green]")
        lines.append(f"   VRAM: [bold]{hw.vram_gb} GB[/bold]    "
                     f"RAM: [bold]{hw.ram_gb} GB[/bold]")
    elif hw.gpu_type == "amd":
        lines.append(f"{gpu_icon}  [bold bright_red]AMD {hw.gpu_name}[/bold bright_red]")
        lines.append(f"   RAM: [bold]{hw.ram_gb} GB[/bold]")
    else:
        lines.append(f"{gpu_icon}  [dim]CPU only[/dim]   RAM: [bold]{hw.ram_gb} GB[/bold]")

    if hw.recommended:
        lines.append("")
        lines.append("[bold white]Recommended for this machine:[/bold white]")
        for mid in hw.recommended:
            m = MODEL_LOOKUP.get(mid, {})
            star = " [bold yellow]★[/bold yellow]" if m.get("recommended") else "  "
            lines.append(
                f"  •{star} [yellow]{mid:<26}[/yellow]"
                f"[dim]{m.get('desc', '')}[/dim]"
            )
    else:
        lines.append("[dim]No suitable models found — try a smaller model.[/dim]")

    body = "\n".join(lines)
    console.print(
        Panel(
            body,
            title="[bold cyan]Hardware Detected[/bold cyan]",
            border_style="cyan",
            padding=(1, 2),
        )
    )


# ── Visual helpers ────────────────────────────────────────────────────────────

BANNER_LINES = [
    ("bold bright_cyan",
     "  ██╗    ██╗ ██████╗ ██████╗ ██╗  ██╗███████╗ █████╗ ███████╗███████╗   █████╗ ██╗"),
    ("bold cyan",
     "  ██║    ██║██╔═══██╗██╔══██╗██║ ██╔╝██╔════╝██╔══██╗██╔════╝██╔════╝  ██╔══██╗██║"),
    ("bold blue",
     "  ██║ █╗ ██║██║   ██║██████╔╝█████╔╝ ███████╗███████║█████╗  █████╗    ███████║██║"),
    ("bold bright_blue",
     "  ██║███╗██║██║   ██║██╔══██╗██╔═██╗ ╚════██║██╔══██║██╔══╝  ██╔══╝    ██╔══██║██║"),
    ("bold magenta",
     "  ╚███╔███╔╝╚██████╔╝██║  ██║██║  ██╗███████║██║  ██║██║     ███████╗  ██║  ██║██║"),
    ("bold bright_magenta",
     "   ╚══╝╚══╝  ╚═════╝ ╚═╝  ╚═╝╚═╝  ╚═╝╚══════╝╚═╝  ╚═╝╚═╝     ╚══════╝  ╚═╝  ╚═╝╚═╝"),
]


def print_banner() -> None:
    console.print()
    for style, line in BANNER_LINES:
        console.print(line, style=style)
    console.print()

    grid = Table.grid(expand=True)
    grid.add_column(justify="center")
    grid.add_row(Text(
        "🔒  Private AI · 100% Offline · Your Data Never Leaves This Machine",
        style="bold white",
    ))
    grid.add_row(Text(
        f"v{APP_VERSION}  ·  Powered by Ollama  ·  GPL-3.0  ·  No API Keys  ·  No Tracking",
        style="dim white",
    ))
    console.print(Panel(grid, border_style="bright_cyan", padding=(0, 2)))
    console.print()


def print_help() -> None:
    t = Table(
        title="Commands", box=box.ROUNDED,
        border_style="cyan", title_style="bold cyan",
    )
    t.add_column("Command",     style="bold yellow", no_wrap=True)
    t.add_column("Description", style="white")
    rows = [
        ("/help",                    "Show this help message"),
        ("/models",                  "Browse and switch to a different model"),
        ("/new",                     "Start a fresh conversation"),
        ("/history",                 "Print conversation history"),
        ("/search <term>",           "Search conversation for a keyword"),
        ("/export",                  "Export as Markdown (default)"),
        ("/export html|pdf|docx|json","Export in a specific format"),
        ("/export <path>",           "Export to path — format inferred from extension"),
        ("/image <path>",            "Attach image for next message (vision models)"),
        ("/session",                 "List saved sessions"),
        ("/session save [name]",     "Save current conversation"),
        ("/session load <name>",     "Restore a saved conversation"),
        ("/session delete <name>",   "Delete a saved session"),
        ("/system",                  "Show current system prompt"),
        ("/system reset",            "Reset system prompt to default"),
        ("/system <text>",           "Set a custom system prompt"),
        ("/clear",                   "Clear the screen"),
        ("/about",                   "Privacy and licence information"),
        ("/quit  /exit",             "Exit Worksafe AI"),
    ]
    for cmd, desc in rows:
        t.add_row(cmd, desc)
    console.print(Padding(t, (1, 0)))


def print_about() -> None:
    text = (
        "## 🔒 About Worksafe AI\n\n"
        "Run a fully **private AI assistant** on your own hardware — no cloud, no tracking.\n\n"
        "### Why local?\n"
        "- **Company-policy safe** — confidential data never leaves your machine\n"
        "- **No API keys** or subscriptions required\n"
        "- **Works offline** — no internet needed after first model download\n"
        "- **GDPR / HIPAA / NDA friendly** — complete data sovereignty\n\n"
        "### How it works\n"
        "[Ollama](https://ollama.ai) runs open-weight models locally via a REST API on "
        "`localhost:11434`. This script is a client to that API — nothing more.\n\n"
        "### Privacy guarantee\n"
        "Every prompt stays on **this machine**. No telemetry. No logs sent anywhere.\n\n"
        "### Licence\n"
        "Worksafe AI is free software — [GNU GPL v3](https://www.gnu.org/licenses/gpl-3.0.html)."
    )
    console.print(Panel(
        Markdown(text),
        title="[bold cyan]About Worksafe AI[/bold cyan]",
        border_style="cyan", padding=(1, 2),
    ))


# ── Ollama helpers ────────────────────────────────────────────────────────────

def ollama_running() -> bool:
    try:
        r = requests.get(f"{OLLAMA_API}/api/tags", timeout=3)
        return r.status_code == 200
    except requests.exceptions.ConnectionError:
        return False


def start_ollama() -> bool:
    if shutil.which("ollama") is None:
        return False
    console.print("[yellow]Starting Ollama daemon…[/yellow]")
    subprocess.Popen(
        ["ollama", "serve"],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        start_new_session=True,
    )
    for _ in range(15):
        time.sleep(1)
        if ollama_running():
            console.print("[green]✓ Ollama is running[/green]")
            return True
    return False


def installed_models() -> list[str]:
    try:
        r = requests.get(f"{OLLAMA_API}/api/tags", timeout=5)
        return [m["name"] for m in r.json().get("models", [])]
    except Exception:
        return []


def pull_model(model: str) -> None:
    console.print(f"\n[bold cyan]Downloading [yellow]{model}[/yellow] — this may take a while…[/bold cyan]\n")
    with requests.post(
        f"{OLLAMA_API}/api/pull", json={"name": model}, stream=True, timeout=None
    ) as resp:
        last_pct = -1
        for raw in resp.iter_lines():
            if not raw:
                continue
            try:
                d = json.loads(raw)
            except json.JSONDecodeError:
                continue
            status = d.get("status", "")
            total  = d.get("total", 0)
            comp   = d.get("completed", 0)
            if total and comp:
                pct = int(comp / total * 100)
                if pct != last_pct and pct % 5 == 0:
                    bar = "█" * (pct // 5) + "░" * (20 - pct // 5)
                    console.print(
                        f"  [cyan][{bar}][/cyan] [bold]{pct:3d}%[/bold]  [dim]{status}[/dim]",
                        end="\r",
                    )
                    last_pct = pct
            elif status:
                console.print(f"  [dim]{status:<60}[/dim]", end="\r")
    console.print(f"\n[bold green]✓ {model} downloaded successfully![/bold green]\n")


# ── Model selection ───────────────────────────────────────────────────────────

def choose_model(
    preselected: Optional[str] = None,
    hw: Optional[HardwareInfo] = None,
) -> str:
    local = installed_models()

    console.print()
    console.print(Rule("[bold cyan]Model Selection[/bold cyan]", style="cyan"))
    console.print()

    # Group by category
    categories: dict[str, list[dict]] = {}
    for m in MODELS:
        categories.setdefault(m["category"], []).append(m)

    counter = 1
    index_map: dict[int, str] = {}

    for cat, models in categories.items():
        t = Table(
            title=cat, box=box.SIMPLE_HEAD,
            border_style="dim blue", title_style="bold white",
            show_header=True, header_style="bold dim white",
        )
        t.add_column("#",       width=4,  justify="right", style="bold white")
        t.add_column("Model",   min_width=26, no_wrap=True)
        t.add_column("Description", min_width=38)
        t.add_column("Size",    width=9,  justify="center", style="dim cyan")
        t.add_column("Min RAM", width=9,  justify="center", style="dim magenta")
        t.add_column("Status",  width=14, justify="center")

        for m in models:
            is_local = any(m["id"].split(":")[0] in loc for loc in local)
            is_rec   = m.get("recommended") or (hw and m["id"] in hw.recommended)
            badge    = (
                "[bold green]● Installed[/bold green]"
                if is_local else "[dim]○ Available[/dim]"
            )
            star = " [bold yellow]★[/bold yellow]" if is_rec else ""
            t.add_row(
                str(counter),
                f"[{m['color']}]{m['id']}[/{m['color']}]{star}",
                f"{m['desc']}",
                m["size"],
                f"{m['min_ram']} GB",
                badge,
            )
            index_map[counter] = m["id"]
            counter += 1

        console.print(t)

    console.print(
        "[dim]  ★ = recommended for your hardware    "
        "● = already downloaded[/dim]\n"
    )

    if preselected and _is_installed(preselected, local):
        return preselected

    while True:
        choice = Prompt.ask(
            "[bold cyan]Select a model[/bold cyan] [dim](number or full name)[/dim]"
        )
        model: Optional[str] = None
        if choice.isdigit() and int(choice) in index_map:
            model = index_map[int(choice)]
        elif choice in MODEL_LOOKUP:
            model = choice
        elif _is_installed(choice, local):
            model = choice
        else:
            console.print("[red]Invalid selection — enter a number or a model name[/red]")
            continue

        if not _is_installed(model, local):
            if Confirm.ask(
                f"\n  [yellow]{model}[/yellow] is not downloaded yet. Download it now?",
                default=True,
            ):
                pull_model(model)
            else:
                console.print("[red]Please select an already-installed model or download one.[/red]")
                continue
        return model


# ── Export conversation ───────────────────────────────────────────────────────

def _export_markdown(messages: list[dict], model: str, filepath: Path) -> Path:
    ts          = datetime.now()
    model_short = model.split(":")[0].capitalize()
    user_turns  = [m for m in messages if m["role"] == "user"]

    lines = [
        "# Worksafe AI — Conversation Export", "",
        "| Field    | Value |",
        "|----------|-------|",
        f"| Date     | {ts.strftime('%Y-%m-%d %H:%M:%S')} |",
        f"| Model    | `{model}` |",
        f"| Turns    | {len(user_turns)} |",
        f"| Version  | Worksafe AI v{APP_VERSION} |",
        "", "---", "",
    ]
    for msg in messages:
        role, content = msg["role"], msg["content"].strip()
        if role == "system":
            lines += [f"> **System Prompt:** {content}", "", "---", ""]
        elif role == "user":
            lines += [f"### 🧑 You", "", content, ""]
        elif role == "assistant":
            lines += [f"### 🤖 {model_short}", "", content, "", "---", ""]
    lines += [
        "",
        "*Generated by [Worksafe AI](https://github.com/hardlygospel/worksafe-ai)"
        " — 100% private, running on your own hardware.*",
    ]
    filepath.write_text("\n".join(lines), encoding="utf-8")
    return filepath


def _export_html(messages: list[dict], model: str, filepath: Path) -> Path:
    ts          = datetime.now()
    model_short = model.split(":")[0].capitalize()
    user_turns  = [m for m in messages if m["role"] == "user"]

    def esc(s: str) -> str:
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    body_parts: list[str] = []
    for msg in messages:
        role, content = msg["role"], msg["content"].strip()
        if role == "system":
            body_parts.append(
                f'<div class="system"><strong>System Prompt</strong><br>{esc(content)}</div>'
            )
        elif role == "user":
            body_parts.append(
                f'<div class="bubble user"><span class="label">You</span>'
                f'<p>{esc(content)}</p></div>'
            )
        elif role == "assistant":
            body_parts.append(
                f'<div class="bubble ai"><span class="label">{esc(model_short)}</span>'
                f'<p>{esc(content)}</p></div>'
            )

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Worksafe AI Export — {ts.strftime('%Y-%m-%d')}</title>
<style>
  :root {{
    --bg:#1a1b1e; --surface:#25262b; --border:#373a40;
    --text:#c1c2c5; --dim:#5c5f66;
    --user:#1971c2; --ai:#2c2e33; --system:#1c3a1c;
    --accent:#339af0; --green:#40c057;
  }}
  @media(prefers-color-scheme:light) {{
    :root {{
      --bg:#f8f9fa; --surface:#fff; --border:#dee2e6;
      --text:#212529; --dim:#868e96;
      --user:#1971c2; --ai:#f1f3f5; --system:#e9f5e9;
    }}
  }}
  *{{box-sizing:border-box; margin:0; padding:0}}
  body{{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;
    background:var(--bg); color:var(--text); padding:2rem 1rem; line-height:1.6}}
  .wrap{{max-width:820px; margin:0 auto}}
  header{{text-align:center; margin-bottom:2rem; padding-bottom:1.5rem;
    border-bottom:1px solid var(--border)}}
  header h1{{font-size:1.8rem; color:var(--accent); margin-bottom:.4rem}}
  header p{{color:var(--dim); font-size:.85rem}}
  .meta{{display:inline-flex; gap:1.5rem; margin-top:.8rem; font-size:.8rem;
    color:var(--dim)}}
  .meta span b{{color:var(--text)}}
  .system{{background:var(--system); border:1px solid var(--border);
    border-radius:8px; padding:.8rem 1rem; margin:1rem 0;
    font-size:.82rem; color:var(--dim)}}
  .bubble{{margin:.8rem 0; padding:1rem 1.2rem; border-radius:12px;
    border:1px solid var(--border)}}
  .bubble.user{{background:var(--user); color:#fff; border-color:transparent;
    margin-left:8%}}
  .bubble.ai{{background:var(--ai); margin-right:8%}}
  .label{{font-size:.75rem; font-weight:700; text-transform:uppercase;
    letter-spacing:.05em; opacity:.7; display:block; margin-bottom:.4rem}}
  .bubble.user .label{{color:rgba(255,255,255,.7)}}
  .bubble.ai .label{{color:var(--green)}}
  p{{white-space:pre-wrap; word-break:break-word}}
  footer{{text-align:center; margin-top:2.5rem; padding-top:1rem;
    border-top:1px solid var(--border); font-size:.78rem; color:var(--dim)}}
  footer a{{color:var(--accent); text-decoration:none}}
</style>
</head>
<body>
<div class="wrap">
  <header>
    <h1>🔒 Worksafe AI</h1>
    <p>Conversation Export</p>
    <div class="meta">
      <span><b>Date</b> {ts.strftime('%Y-%m-%d %H:%M')}</span>
      <span><b>Model</b> {esc(model)}</span>
      <span><b>Turns</b> {len(user_turns)}</span>
    </div>
  </header>
  {''.join(body_parts)}
  <footer>
    Generated by <a href="https://github.com/hardlygospel/worksafe-ai">Worksafe AI</a>
    v{APP_VERSION} — 100% private, running on your own hardware.
  </footer>
</div>
</body>
</html>"""

    filepath.write_text(html, encoding="utf-8")
    return filepath


def _ensure_fpdf2() -> None:
    try:
        import fpdf  # noqa: F401
    except ImportError:
        console.print("[yellow]Installing fpdf2 for PDF export…[/yellow]")
        _pip_install(["fpdf2>=2.7.0"])


def _export_pdf(messages: list[dict], model: str, filepath: Path) -> Path:
    _ensure_fpdf2()
    from fpdf import FPDF  # noqa: PLC0415

    ts          = datetime.now()
    model_short = model.split(":")[0].capitalize()
    user_turns  = [m for m in messages if m["role"] == "user"]

    def safe(s: str) -> str:
        return s.encode("latin-1", errors="replace").decode("latin-1")

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(0, 120, 190)
    pdf.cell(0, 12, "Worksafe AI", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(110, 110, 110)
    pdf.cell(0, 8, "Conversation Export", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(150, 150, 150)
    meta = (f"Date: {ts.strftime('%Y-%m-%d %H:%M:%S')}   "
            f"Model: {model}   Turns: {len(user_turns)}")
    pdf.cell(0, 6, safe(meta), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_draw_color(210, 210, 210)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(6)

    for msg in messages:
        role, content = msg["role"], msg["content"].strip()
        if role == "system":
            pdf.set_font("Helvetica", "I", 8)
            pdf.set_text_color(160, 160, 160)
            excerpt = content[:300] + ("…" if len(content) > 300 else "")
            pdf.multi_cell(0, 5, safe(f"System: {excerpt}"))
            pdf.ln(4)
        elif role == "user":
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(20, 90, 190)
            pdf.cell(0, 7, "You", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(45, 45, 45)
            pdf.multi_cell(0, 6, safe(content))
            pdf.ln(3)
        elif role == "assistant":
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(10, 130, 70)
            pdf.cell(0, 7, safe(model_short), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(45, 45, 45)
            pdf.multi_cell(0, 6, safe(content))
            pdf.ln(3)
            pdf.set_draw_color(225, 225, 225)
            pdf.line(15, pdf.get_y(), 195, pdf.get_y())
            pdf.ln(5)

    pdf.ln(4)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(160, 160, 160)
    pdf.cell(
        0, 6,
        safe(f"Generated by Worksafe AI v{APP_VERSION} — 100% private, local AI."),
        align="C",
    )

    pdf.output(str(filepath))
    return filepath


def _ensure_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        console.print("[yellow]Installing python-docx for DOCX export…[/yellow]")
        _pip_install(["python-docx>=1.0.0"])


def _export_docx(messages: list[dict], model: str, filepath: Path) -> Path:
    _ensure_docx()
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt, RGBColor  # noqa: PLC0415
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: PLC0415

    ts          = datetime.now()
    model_short = model.split(":")[0].capitalize()
    user_turns  = [m for m in messages if m["role"] == "user"]

    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Worksafe AI — Conversation Export")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x00, 0x78, 0xBE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Date: {ts.strftime('%Y-%m-%d %H:%M')}   "
        f"Model: {model}   Turns: {len(user_turns)}"
    )
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_paragraph()

    # ── Messages ───────────────────────────────────────────────────────────────
    for msg in messages:
        role, content = msg["role"], msg["content"].strip()
        if role == "system":
            p = doc.add_paragraph()
            r = p.add_run(f"System: {content}")
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            doc.add_paragraph()
        elif role == "user":
            lp = doc.add_paragraph()
            lr = lp.add_run("You")
            lr.bold = True
            lr.font.color.rgb = RGBColor(0x14, 0x5A, 0xBE)
            doc.add_paragraph(content)
        elif role == "assistant":
            lp = doc.add_paragraph()
            lr = lp.add_run(model_short)
            lr.bold = True
            lr.font.color.rgb = RGBColor(0x0A, 0x82, 0x46)
            doc.add_paragraph(content)
            doc.add_paragraph()

    # ── Footer ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fr = fp.add_run(f"Generated by Worksafe AI v{APP_VERSION} — 100% private, local AI.")
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(str(filepath))
    return filepath


def _export_json(messages: list[dict], model: str, filepath: Path) -> Path:
    ts   = datetime.now()
    data = {
        "meta": {
            "app":        "Worksafe AI",
            "version":    APP_VERSION,
            "exported_at": ts.isoformat(),
            "model":      model,
            "turn_count": len([m for m in messages if m["role"] == "user"]),
        },
        "messages": [
            {"index": i, "role": m["role"], "content": m["content"]}
            for i, m in enumerate(messages)
        ],
    }
    filepath.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    return filepath


def export_conversation(
    messages: list[dict],
    model: str,
    filepath: Optional[Path] = None,
    fmt: str = "md",
) -> Path:
    ts_str = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    ext    = {"md": ".md", "html": ".html", "pdf": ".pdf",
              "docx": ".docx", "json": ".json"}.get(fmt, ".md")
    if filepath is None:
        filepath = Path.cwd() / f"worksafe-ai-{ts_str}{ext}"

    dispatch = {
        "html": _export_html,
        "pdf":  _export_pdf,
        "docx": _export_docx,
        "json": _export_json,
    }
    return dispatch.get(fmt, _export_markdown)(messages, model, filepath)


# ── Session management ────────────────────────────────────────────────────────

def _session_path(name: str) -> Path:
    safe = re.sub(r"[^\w\-]", "_", name.strip())
    return SESSIONS_DIR / f"{safe}.json"


def save_session(name: str, model: str, messages: list[dict], system_prompt: str) -> Path:
    SESSIONS_DIR.mkdir(parents=True, exist_ok=True)
    path = _session_path(name)
    data = {
        "name":          name,
        "model":         model,
        "system_prompt": system_prompt,
        "version":       APP_VERSION,
        "saved_at":      datetime.now().isoformat(),
        "messages":      messages,
    }
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    return path


def load_session(name: str) -> dict:
    path = _session_path(name)
    if not path.exists():
        # Try partial match
        matches = list(SESSIONS_DIR.glob(f"*{re.sub(r'[^\\w]', '_', name)}*.json"))
        if len(matches) == 1:
            path = matches[0]
        else:
            raise FileNotFoundError(f"Session '{name}' not found.")
    return json.loads(path.read_text(encoding="utf-8"))


def list_sessions() -> list[dict]:
    if not SESSIONS_DIR.exists():
        return []
    sessions = []
    for f in sorted(SESSIONS_DIR.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True):
        try:
            d = json.loads(f.read_text(encoding="utf-8"))
            sessions.append({
                "name":       d.get("name", f.stem),
                "model":      d.get("model", "unknown"),
                "saved_at":   d.get("saved_at", "")[:16].replace("T", " "),
                "turns":      len([m for m in d.get("messages", []) if m["role"] == "user"]),
                "file":       f,
            })
        except Exception:
            pass
    return sessions


def delete_session(name: str) -> bool:
    path = _session_path(name)
    if path.exists():
        path.unlink()
        return True
    return False


# ── Chat engine ───────────────────────────────────────────────────────────────

def stream_chat(model: str, messages: list[dict]) -> str:
    full = ""
    try:
        with requests.post(
            f"{OLLAMA_API}/api/chat",
            json={"model": model, "messages": messages, "stream": True},
            stream=True,
            timeout=180,
        ) as resp:
            resp.raise_for_status()
            for raw in resp.iter_lines():
                if not raw:
                    continue
                try:
                    data = json.loads(raw)
                except json.JSONDecodeError:
                    continue
                token = data.get("message", {}).get("content", "")
                if token:
                    full += token
                    console.print(token, end="", markup=False, highlight=False)
                if data.get("done"):
                    break
    except requests.exceptions.RequestException as e:
        console.print(f"\n[red]Connection error: {e}[/red]")
    return full


def chat_loop(model: str, initial_system: Optional[str] = None) -> None:
    current_system  = initial_system or DEFAULT_SYSTEM_PROMPT
    messages: list[dict] = [{"role": "system", "content": current_system}]
    model_short     = model.split(":")[0].capitalize()
    pending_image:  Optional[Path] = None   # image attached for the next message

    console.print()
    console.print(Panel(
        f"[bold green]✓  Connected to [yellow]{model}[/yellow][/bold green]\n"
        "[dim]Type a message and press Enter.  [bold]/help[/bold] for commands.[/dim]",
        border_style="green", padding=(0, 2),
    ))
    console.print()

    while True:
        prompt_suffix = (
            f" [dim cyan]📎 {pending_image.name}[/dim cyan]"
            if pending_image else ""
        )
        try:
            user_input = Prompt.ask(
                f"[bold bright_green]You[/bold bright_green]{prompt_suffix}"
            ).strip()
        except (KeyboardInterrupt, EOFError):
            console.print("\n[dim]Session ended. All data stayed on this machine.[/dim]")
            break

        if not user_input:
            continue

        cmd     = user_input.lower().strip()
        cmd_raw = user_input.strip()

        # ── Quit ──────────────────────────────────────────────────────────────
        if cmd in ("/quit", "/exit"):
            console.print("[dim]Goodbye! Your session was completely private.[/dim]")
            break

        # ── Help ──────────────────────────────────────────────────────────────
        if cmd == "/help":
            print_help()
            continue

        # ── About ─────────────────────────────────────────────────────────────
        if cmd == "/about":
            print_about()
            continue

        # ── Clear ─────────────────────────────────────────────────────────────
        if cmd == "/clear":
            console.clear()
            print_banner()
            continue

        # ── New conversation ──────────────────────────────────────────────────
        if cmd == "/new":
            messages      = [{"role": "system", "content": current_system}]
            pending_image = None
            console.print("[yellow]Conversation cleared — fresh start.[/yellow]")
            continue

        # ── History ───────────────────────────────────────────────────────────
        if cmd == "/history":
            turns = [m for m in messages if m["role"] != "system"]
            if not turns:
                console.print("[dim]No conversation history yet.[/dim]")
            else:
                for m in turns:
                    style = "bright_green" if m["role"] == "user" else "bright_cyan"
                    label = "You" if m["role"] == "user" else model_short
                    console.print(f"[{style}]{label}:[/{style}] {m['content']}\n")
            continue

        # ── Search ────────────────────────────────────────────────────────────
        if cmd.startswith("/search"):
            term = cmd_raw[7:].strip()
            if not term:
                console.print("[yellow]Usage: /search <term>[/yellow]")
                continue
            hits = [
                (i, m) for i, m in enumerate(messages)
                if m["role"] in ("user", "assistant")
                and term.lower() in m["content"].lower()
            ]
            if not hits:
                console.print(f"[dim]No matches for '[yellow]{term}[/yellow]'[/dim]")
            else:
                console.print(
                    f"\n[bold cyan]Found [yellow]{len(hits)}[/yellow] "
                    f"match(es) for '[yellow]{term}[/yellow]':[/bold cyan]\n"
                )
                for idx, m in hits:
                    role_style = "bright_green" if m["role"] == "user" else "bright_cyan"
                    label      = "You" if m["role"] == "user" else model_short
                    content    = m["content"]
                    pos        = content.lower().find(term.lower())
                    ctx_start  = max(0, pos - 80)
                    ctx_end    = min(len(content), pos + len(term) + 80)
                    excerpt    = (
                        ("…" if ctx_start > 0 else "")
                        + content[ctx_start:ctx_end]
                        + ("…" if ctx_end < len(content) else "")
                    )
                    # Rich Text for safe highlighting
                    snippet = Text(excerpt)
                    snippet.highlight_words([term], style="bold yellow", case_sensitive=False)
                    console.print(
                        f"  [{role_style}]{label}[/{role_style}] "
                        f"[dim](message #{idx})[/dim]"
                    )
                    console.print("  ", snippet, "\n")
            continue

        # ── Export ────────────────────────────────────────────────────────────
        if cmd.startswith("/export"):
            turns = [m for m in messages if m["role"] != "system"]
            if not turns:
                console.print("[yellow]Nothing to export — have a conversation first.[/yellow]")
                continue
            arg      = cmd_raw[7:].strip()
            fmt      = "md"
            out_path: Optional[Path] = None
            if arg:
                key = arg.lower()
                if key in ("html", "htm"):
                    fmt = "html"
                elif key == "pdf":
                    fmt = "pdf"
                elif key in ("docx", "word"):
                    fmt = "docx"
                elif key == "json":
                    fmt = "json"
                elif key in ("md", "markdown"):
                    fmt = "md"
                else:
                    out_path = Path(arg).expanduser()
                    fmt = {
                        ".html": "html", ".htm": "html",
                        ".pdf":  "pdf",
                        ".docx": "docx",
                        ".json": "json",
                    }.get(out_path.suffix.lower(), "md")
            saved = export_conversation(messages, model, out_path, fmt)
            fmt_label = {
                "md": "Markdown", "html": "HTML",
                "pdf": "PDF", "docx": "Word (DOCX)", "json": "JSON",
            }[fmt]
            console.print(Panel(
                f"[bold green]✓  Exported as {fmt_label}:[/bold green]\n"
                f"[cyan]{saved}[/cyan]\n\n"
                f"[dim]{len(turns)} message(s) · model: {model}[/dim]",
                title="[bold]Export[/bold]", border_style="green", padding=(0, 2),
            ))
            continue

        # ── Image attachment ──────────────────────────────────────────────────
        if cmd.startswith("/image"):
            if not _is_vision_model(model):
                console.print(Panel(
                    "[yellow]The current model doesn't support images.[/yellow]\n"
                    "Switch to a vision model first:\n"
                    "  [cyan]llava:7b  llava:13b  llava-llama3:8b  moondream:1.8b[/cyan]",
                    border_style="yellow", padding=(0, 2),
                ))
                continue
            img_arg = cmd_raw[6:].strip()
            if not img_arg:
                console.print("[yellow]Usage: /image <path>[/yellow]")
                continue
            img_path = Path(img_arg).expanduser()
            if not img_path.exists():
                console.print(f"[red]File not found: {img_path}[/red]")
                continue
            if img_path.suffix.lower() not in {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"}:
                console.print("[red]Unsupported image type. Use: jpg, png, gif, webp, bmp[/red]")
                continue
            pending_image = img_path
            console.print(
                f"[green]✓  Image ready: [cyan]{img_path.name}[/cyan]. "
                f"Now type your question.[/green]"
            )
            continue

        # ── Sessions ──────────────────────────────────────────────────────────
        if cmd.startswith("/session"):
            arg = cmd_raw[8:].strip()
            sub = arg.split(None, 1)
            subcmd = sub[0].lower() if sub else ""

            if not subcmd or subcmd == "list":
                sessions = list_sessions()
                if not sessions:
                    console.print("[dim]No saved sessions yet. Use [bold]/session save[/bold] to create one.[/dim]")
                else:
                    t = Table(title="Saved Sessions", box=box.ROUNDED,
                              border_style="cyan", title_style="bold cyan")
                    t.add_column("Name",    style="bold yellow", min_width=20)
                    t.add_column("Model",   style="cyan",        min_width=16)
                    t.add_column("Turns",   style="white",       width=7,  justify="right")
                    t.add_column("Saved",   style="dim",         min_width=16)
                    for s in sessions:
                        t.add_row(s["name"], s["model"], str(s["turns"]), s["saved_at"])
                    console.print(Padding(t, (1, 0)))

            elif subcmd == "save":
                name = sub[1].strip() if len(sub) > 1 else ""
                if not name:
                    try:
                        name = Prompt.ask("[cyan]Session name[/cyan]").strip()
                    except (KeyboardInterrupt, EOFError):
                        console.print("[dim]Save cancelled.[/dim]")
                        continue
                if not name:
                    console.print("[red]Session name cannot be empty.[/red]")
                    continue
                path = save_session(name, model, messages, current_system)
                console.print(f"[green]✓  Session '[yellow]{name}[/yellow]' saved.[/green]")

            elif subcmd == "load":
                name = sub[1].strip() if len(sub) > 1 else ""
                if not name:
                    console.print("[yellow]Usage: /session load <name>[/yellow]")
                    continue
                try:
                    data          = load_session(name)
                    model         = data["model"]
                    model_short   = model.split(":")[0].capitalize()
                    current_system= data.get("system_prompt", DEFAULT_SYSTEM_PROMPT)
                    messages      = data["messages"]
                    pending_image = None
                    console.print(Panel(
                        f"[bold green]✓  Session '[yellow]{data['name']}[/yellow]' restored.[/bold green]\n"
                        f"Model: [yellow]{model}[/yellow]   "
                        f"Turns: [white]{len([m for m in messages if m['role'] == 'user'])}[/white]",
                        border_style="green", padding=(0, 2),
                    ))
                except FileNotFoundError as e:
                    console.print(f"[red]{e}[/red]")

            elif subcmd == "delete":
                name = sub[1].strip() if len(sub) > 1 else ""
                if not name:
                    console.print("[yellow]Usage: /session delete <name>[/yellow]")
                    continue
                if delete_session(name):
                    console.print(f"[green]✓  Session '[yellow]{name}[/yellow]' deleted.[/green]")
                else:
                    console.print(f"[red]Session '{name}' not found.[/red]")
            else:
                console.print("[red]Unknown session sub-command. "
                              "Use: list · save [name] · load <name> · delete <name>[/red]")
            continue

        # ── System prompt ─────────────────────────────────────────────────────
        if cmd.startswith("/system"):
            arg = cmd_raw[7:].strip()
            if not arg or arg.lower() == "show":
                console.print(Panel(
                    current_system,
                    title="[bold cyan]Current System Prompt[/bold cyan]",
                    border_style="cyan", padding=(0, 2),
                ))
            elif arg.lower() == "reset":
                current_system = DEFAULT_SYSTEM_PROMPT
                messages[0]    = {"role": "system", "content": current_system}
                console.print("[green]✓  System prompt reset to default.[/green]")
            else:
                current_system = arg
                messages[0]    = {"role": "system", "content": current_system}
                console.print(Panel(
                    f"[bold green]✓  System prompt updated:[/bold green]\n{current_system}",
                    border_style="green", padding=(0, 2),
                ))
            continue

        # ── Switch model ──────────────────────────────────────────────────────
        if cmd == "/models":
            model         = choose_model(model)
            model_short   = model.split(":")[0].capitalize()
            messages      = [{"role": "system", "content": current_system}]
            pending_image = None   # clear image if model no longer supports it
            console.print(
                f"[green]Switched to [yellow]{model}[/yellow]. Conversation reset.[/green]"
            )
            continue

        # ── Send to model ─────────────────────────────────────────────────────
        user_msg: dict = {"role": "user", "content": user_input}
        if pending_image and _is_vision_model(model):
            try:
                b64 = base64.b64encode(pending_image.read_bytes()).decode()
                user_msg["images"] = [b64]
                console.print(
                    f"[dim cyan]📎 Sending with image: {pending_image.name}[/dim cyan]"
                )
            except OSError as e:
                console.print(f"[red]Could not read image: {e}[/red]")
            pending_image = None

        messages.append(user_msg)
        console.print()
        console.print(f"[bold bright_cyan]{model_short}[/bold bright_cyan] ", end="")
        response = stream_chat(model, messages)
        console.print("\n")
        console.print(Rule(style="dim blue"))
        console.print()
        if response:
            messages.append({"role": "assistant", "content": response})


# ── Entry point ───────────────────────────────────────────────────────────────

def _parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        prog="worksafe_ai",
        description="Worksafe AI — Private local LLM. 100%% offline. GPL-3.0.",
    )
    p.add_argument("-m", "--model",     help="Skip model selection and use this model directly")
    p.add_argument("-s", "--system",    help="Custom system prompt (overrides default)")
    p.add_argument("--no-gpu-check",   action="store_true", help="Skip hardware detection")
    p.add_argument("--version",        action="version", version=f"Worksafe AI v{APP_VERSION}")
    return p.parse_args()


def main() -> None:
    args = _parse_args()
    print_banner()

    # 1. Hardware detection
    hw: Optional[HardwareInfo] = None
    if not args.no_gpu_check:
        hw = detect_hardware()
        show_hardware_panel(hw)
        console.print()

    # 2. Check Ollama binary
    if shutil.which("ollama") is None:
        console.print(Panel(
            "[bold red]Ollama is not installed.[/bold red]\n\n"
            "Install it with one of:\n"
            "  [bold yellow]macOS / Linux:[/bold yellow]   "
            "[cyan]curl -fsSL https://ollama.ai/install.sh | sh[/cyan]\n"
            "  [bold yellow]macOS Homebrew:[/bold yellow]  "
            "[cyan]brew install ollama[/cyan]\n"
            "  [bold yellow]Windows:[/bold yellow]         "
            "[cyan]https://ollama.ai/download[/cyan]\n\n"
            "Then re-run this script.",
            title="[red]Installation Required[/red]",
            border_style="red", padding=(1, 2),
        ))
        sys.exit(1)

    # 3. Ensure daemon is running
    if not ollama_running():
        if not start_ollama():
            console.print(
                "[bold red]Could not start Ollama. "
                "Run `ollama serve` in another terminal then retry.[/bold red]"
            )
            sys.exit(1)
    else:
        console.print("[green]✓  Ollama is running[/green]\n")

    # 4. Model selection
    model = choose_model(preselected=args.model, hw=hw)

    # 5. Chat
    chat_loop(model, initial_system=args.system)


if __name__ == "__main__":
    main()
